"""DOCX parser using python-docx."""

import logging
from pathlib import Path

from services.documents.parsers.pdf_parser import ParsedDocument, ParsedPage

logger = logging.getLogger(__name__)


def parse_docx(file_path: str | Path, file_name: str | None = None) -> ParsedDocument:
    """Extract text from DOCX files.

    DOCX files don't have true page numbers (they're flow-based),
    so we assign page_number=1 for all content.
    """
    from docx import Document

    path = Path(file_path)
    fname = file_name or path.name
    doc = Document(str(path))

    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    full_text = "\n\n".join(paragraphs)

    pages = []
    if full_text.strip():
        pages.append(ParsedPage(page_number=1, text=full_text))

    logger.info(f"Parsed DOCX: {fname}, {len(paragraphs)} paragraphs")

    return ParsedDocument(
        file_name=fname,
        pages=pages,
        total_pages=1,
        metadata={"paragraph_count": len(paragraphs)},
    )
